from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"F:\Internship\Ubuntu\ubuntu-terminal-online")
ASSETS = ROOT / "docs_assets"
OUT_DIR = ROOT / "documentation"
OUT_PATH = OUT_DIR / "Ubuntu_Terminal_Online_Project_Documentation_Faiz_Shaikh.docx"

ORANGE = "E95420"
AUBERGINE = "2C001E"
DARK = "1F2933"
MUTED = "5F6B76"
PALE = "FFF4EF"
LIGHT = "F2F4F7"
WHITE = "FFFFFF"
BORDER = "D7DCE2"
GREEN = "15803D"


def set_run_font(run, name="Calibri", size=None, color=DARK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            set_cell_margins(cell)


def set_table_borders(table, color=BORDER, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_table_no_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "nil")
        borders.append(node)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def define_numbering(doc):
    numbering = doc.part.numbering_part.element

    def make_abstract(abstract_id, fmt, text, left=720, hanging=360, font=None):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(left))
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), str(hanging))
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "160")
        spacing.set(qn("w:line"), "280")
        spacing.set(qn("w:lineRule"), "auto")
        ppr.extend([tabs, ind, spacing])
        lvl.extend([start, num_fmt, lvl_text, lvl_jc, suff, ppr])
        if font:
            rpr = OxmlElement("w:rPr")
            rfonts = OxmlElement("w:rFonts")
            rfonts.set(qn("w:ascii"), font)
            rfonts.set(qn("w:hAnsi"), font)
            rpr.append(rfonts)
            lvl.append(rpr)
        abstract.append(lvl)
        return abstract

    def make_num(num_id, abstract_id):
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        aid = OxmlElement("w:abstractNumId")
        aid.set(qn("w:val"), str(abstract_id))
        num.append(aid)
        return num

    # OOXML requires all abstract numbering definitions before concrete
    # instances. Separate decimal instances restart each independent sequence.
    numbering.append(make_abstract(41, "bullet", "•", font="Arial"))
    numbering.append(make_abstract(42, "decimal", "%1."))
    numbering.append(make_num(41, 41))
    numbering.append(make_num(42, 42))
    numbering.append(make_num(43, 42))
    return 41, 42, 43


def add_list_item(doc, text, num_id, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_body(doc, text, bold_lead=None, align=WD_ALIGN_PARAGRAPH.LEFT, after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if bold_lead and text.startswith(bold_lead):
        a = p.add_run(bold_lead)
        set_run_font(a, bold=True)
        b = p.add_run(text[len(bold_lead):])
        set_run_font(b)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_callout(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.10
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), PALE)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), ORANGE)
    borders.append(left)
    p_pr.extend([shd, borders])
    r1 = p.add_run(label.upper() + "  ")
    set_run_font(r1, size=9.5, color=ORANGE, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.5, color=DARK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    if getattr(doc, "_pending_page_break", False):
        p.paragraph_format.page_break_before = True
        doc._pending_page_break = False
    p.paragraph_format.keep_with_next = True
    return p


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.line_spacing = 1.0
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F6F7F9")
    p_pr.append(shd)
    run = p.add_run(text)
    set_run_font(run, name="Consolas", size=9.2, color=AUBERGINE)
    return p


def add_picture(doc, filename, width_inches, alt_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    shape = p.add_run().add_picture(str(ASSETS / filename), width=Inches(width_inches))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", alt_text)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = False
    r = p.add_run(text)
    set_run_font(r, size=9.2, color=MUTED, italic=True)
    return p


def add_page_break(doc):
    # A page-break-before property on the next heading avoids empty break
    # paragraphs and preserves one stable section with consistent margins.
    doc._pending_page_break = True


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = False

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_specs = {
        1: (16, 16, 8, ORANGE),
        2: (13, 12, 6, ORANGE),
        3: (12, 8, 4, AUBERGINE),
    }
    for level, (size, before, after, color) in heading_specs.items():
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    hp = header.paragraphs[0]
    hp.paragraph_format.space_after = Pt(0)
    hp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    r = hp.add_run("UBUNTU TERMINAL ONLINE")
    set_run_font(r, size=8.5, color=MUTED, bold=True)
    r = hp.add_run("\tPROJECT DOCUMENTATION")
    set_run_font(r, size=8.5, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    r = fp.add_run("Faiz Shaikh  |  Page ")
    set_run_font(r, size=8.5, color=MUTED)
    add_field(fp, "PAGE")

def add_info_table(doc, headers, rows, widths, header_fill=LIGHT, font_size=9.6):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, widths)
    set_table_borders(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=font_size, color=AUBERGINE, bold=True)
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            cell = cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=DARK)
    set_table_geometry(table, widths)
    return table


def build_document():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    bullet_id, architecture_decimal_id, usage_decimal_id = define_numbering(doc)

    props = doc.core_properties
    props.title = "Ubuntu Terminal Online - Project Documentation"
    props.subject = "Browser-based Ubuntu terminal and OS visual learning environment"
    props.author = "Faiz Shaikh"
    props.keywords = "Ubuntu, Linux, terminal, operating systems, gcc, education, browser"
    props.comments = "Prepared for academic submission from the verified project codebase and production screenshots."

    # Cover page - editorial cover pattern with an Ubuntu-inspired named color override.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(52)
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PROJECT DOCUMENTATION")
    set_run_font(r, size=10.5, color=ORANGE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Ubuntu Terminal Online")
    set_run_font(r, size=30, color=AUBERGINE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(38)
    r = p.add_run("A Browser-Based Linux Practice and Operating Systems Learning Environment")
    set_run_font(r, size=14.5, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(42)
    r = p.add_run("Enabling Ubuntu terminal practicals on Windows laboratory computers and other browser-capable devices")
    set_run_font(r, size=11.5, color=DARK, italic=True)

    for label, value in [
        ("Submitted by", "Faiz Shaikh"),
        ("Role", "Teaching Assistant"),
        ("Programme", "M.Tech - First Year"),
        ("Submission", "College Project Report"),
        ("Academic Year", "2026-27"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"{label}: ")
        set_run_font(r, size=10.5, color=MUTED, bold=True)
        r = p.add_run(value)
        set_run_font(r, size=10.5, color=DARK)

    add_page_break(doc)

    # Contents.
    add_heading(doc, "Contents", 1)
    add_body(doc, "The report is organised around the educational problem, the implemented solution, the verified features, and the observed results.", after=12)
    contents = [
        ("Abstract", "3"),
        ("1. Problem Statement", "3"),
        ("2. Project Objectives", "3"),
        ("3. Proposed Solution", "4"),
        ("4. System Architecture", "4"),
        ("5. Technology Stack", "4"),
        ("6. Key Features", "5"),
        ("7. Implementation Details", "6"),
        ("8. How to Use the Website", "7"),
        ("9. Website Screenshots and Feature Demonstration", "8"),
        ("10. Testing and Results", "11"),
        ("11. Benefits and Educational Impact", "12"),
        ("12. Limitations and Future Scope", "12"),
        ("13. Conclusion", "12"),
        ("Appendix A - Supported Command Groups", "13"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_no_borders(table)
    for title, page in contents:
        cells = table.add_row().cells
        for cell in cells:
            set_cell_margins(cell, top=30, bottom=30)
        p1 = cells[0].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r = p1.add_run(title)
        set_run_font(r, size=10.8, color=DARK, bold=title.startswith(("Abstract", "Appendix")))
        p2 = cells[1].paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p2.paragraph_format.space_after = Pt(0)
        r = p2.add_run(page)
        set_run_font(r, size=10.8, color=ORANGE, bold=True)
    set_table_geometry(table, [8350, 1010], indent_dxa=0)

    add_page_break(doc)

    # Page 3: context.
    add_heading(doc, "Abstract", 1)
    add_body(doc, "Many computers in the college laboratory run Microsoft Windows and do not provide a native Ubuntu installation. A few systems may also be unavailable or have configuration issues, which reduces the number of machines on which students can complete Linux, shell scripting, C programming, and Operating Systems practicals.")
    add_body(doc, "To address this constraint, Faiz Shaikh - Teaching Assistant and first-year M.Tech student at the same institution - developed Ubuntu Terminal Online. The website reproduces the familiar Ubuntu GNOME Terminal experience in a browser and provides an educational shell, cached filesystem, C practical workflow, process simulation, preloaded laboratory programs, and interactive concept visualisations. Students can therefore practise from Windows computers or other modern browser-capable devices without changing the host operating system.")
    add_callout(doc, "Project outcome", "A single browser-based learning surface now supports routine Ubuntu terminal practice, C and shell exercises, and visual Operating Systems demonstrations while preserving each student's working files in that browser.")

    add_heading(doc, "1. Problem Statement", 1)
    add_body(doc, "Linux-based laboratory exercises assume that every learner has reliable access to Ubuntu. In practice, many laboratory computers contain Windows, some systems may have installation or maintenance issues, and dual boot or virtual-machine setup can consume class time. Students may also want to practise from personal devices where they cannot install another operating system.")
    add_body(doc, "The academic need was therefore not just a terminal-looking interface. The required solution had to let students enter realistic Linux commands, create and edit files, run shell scripts, compile classroom C programs, observe process behaviour, and return to their work after a page reload.")

    add_heading(doc, "2. Project Objectives", 1)
    for text in [
        "Provide an Ubuntu-like terminal that runs inside a web browser.",
        "Support the commands and scripting constructs required for common Linux and Operating Systems practicals.",
        "Allow students to create, edit, compile, and execute files without requiring Ubuntu on the host computer.",
        "Persist student files in browser storage while also allowing a clean reset for the next laboratory session.",
        "Explain difficult OS concepts through step-by-step visual lessons in addition to command-line practice.",
    ]:
        add_list_item(doc, text, bullet_id)

    add_page_break(doc)

    # Page 4: solution and architecture.
    add_heading(doc, "3. Proposed Solution", 1)
    add_body(doc, "Ubuntu Terminal Online is a Next.js web application that combines a client-side Bash-style interpreter with a virtual Linux filesystem and a classroom-focused C execution engine. It presents the prompt, colours, title bar, keyboard workflow, history, and editor interactions expected from an Ubuntu terminal. For tools that require a full operating system, the command ubuntu-vm opens a full WebVM-based Ubuntu environment inside the application.")

    add_heading(doc, "4. System Architecture", 1)
    architecture_steps = [
        "Presentation layer: React components render the Ubuntu-styled terminal, split view, editor, reset controls, and visual lessons.",
        "Command layer: the shell parser interprets commands, variables, control structures, pipelines, redirection, command substitution, arithmetic, and executable files.",
        "Execution and storage layer: built-in commands operate on a virtual filesystem saved in browser localStorage; compiled classroom C programs run through the educational C interpreter and simulated process table.",
        "Content layer: a server route reads teacher-provided UTF-8 practical files from public/terminal-files and merges them into /home/student for learners.",
        "Extension path: ubuntu-vm embeds webvm.io when a student needs apt, Python, networking, or a complete Linux userland.",
    ]
    for text in architecture_steps:
        add_list_item(doc, text, architecture_decimal_id)

    add_heading(doc, "5. Technology Stack", 1)
    add_info_table(
        doc,
        ["Layer", "Technology", "Purpose"],
        [
            ("Framework", "Next.js 16.3, React 19.2", "Routing, server endpoint, rendering, and application structure"),
            ("Interface", "JavaScript/JSX, CSS, Tailwind CSS 4", "Responsive Ubuntu-style terminal and visual learning pages"),
            ("Local engine", "Custom shell, VFS, Mini-C, process table", "Offline classroom commands, files, compilation, and OS simulation"),
            ("Persistence", "Browser localStorage", "Retains the student's virtual home directory across reloads"),
            ("Full Linux option", "WebVM embedded experience", "Access to a broader Ubuntu environment when internet access is available"),
        ],
        [1800, 2600, 4960],
        font_size=9.2,
    )

    add_page_break(doc)

    # Page 5: features.
    add_heading(doc, "6. Key Features", 1)
    feature_items = [
        ("Ubuntu-like terminal interface. ", "GNOME-terminal-inspired colours, title bar, welcome banner, prompt, command history, keyboard input, clear/exit behaviour, and responsive layout."),
        ("Practical shell environment. ", "Supports variables, aliases, environment variables, if/case conditions, for/while/until loops, functions, tests, pipelines, logical operators, globs, redirection, command substitution, and arithmetic expansion."),
        ("Virtual Linux filesystem. ", "Includes familiar directories, permissions, timestamps, file operations, text processing tools, named FIFOs, and a home directory at /home/student."),
        ("Browser-cached work. ", "Files persist in localStorage across page reloads. A daily reset returns shared laboratory computers to a clean state, while Hard Reset and factory-reset provide manual recovery."),
        ("C practical support. ", "Students can edit C source, compile with gcc/cc/g++, create an executable, accept input with scanf, and use classroom functions such as printf, fork, wait, getpid, pipes, files, and System V message queues."),
        ("Shared split terminal. ", "Two side-by-side terminals use the same virtual filesystem, allowing one pane to create or compile a file and the other to inspect or execute it."),
        ("Integrated editor. ", "nano, vi, vim, and gedit commands open an in-browser editor for the selected virtual file."),
        ("Preloaded practicals. ", "Teacher-managed programs are supplied through the terminal-files endpoint. The current repository includes exercises for practicals 6, 7, and 8 plus reader/writer and send/receive examples."),
        ("Interactive OS Visual Lab. ", "Dedicated lessons demonstrate program-to-process transition, fork(), waiting, zombie and orphan states, and IPC through named pipes, message queues, shared memory, and signals."),
        ("Full Ubuntu escape hatch. ", "The ubuntu-vm command opens a WebVM session for tasks that need apt, python3, networking, or a broader toolchain than the cached classroom shell."),
    ]
    for lead, detail in feature_items:
        add_list_item(doc, lead + detail, bullet_id, bold_lead=lead)
    add_callout(doc, "Accessibility advantage", "Because the main experience runs in a browser, the same lesson can be opened from a Windows lab PC, laptop, or other modern device without repartitioning, dual boot, or a local Ubuntu installation.")

    add_page_break(doc)

    # Page 6: implementation.
    add_heading(doc, "7. Implementation Details", 1)
    add_heading(doc, "7.1 Terminal and Session Management", 2)
    add_body(doc, "UbuntuTerminal.jsx manages the prompt, cursor, command history, output blocks, input mode, editor state, split-terminal labels, and WebVM overlay. A session object keeps the current working directory, shell variables, aliases, history, and references to the shared virtual filesystem.")

    add_heading(doc, "7.2 Shell Parser and Interpreter", 2)
    add_body(doc, "interpreter.js tokenises and parses Bash-like input into executable structures. It expands variables and globs, evaluates arithmetic and command substitution, executes conditional and loop nodes, connects pipeline stages, and applies input/output/error redirection. Built-ins are dispatched from commands.js, while virtual executable files are handled by the execution layer.")

    add_heading(doc, "7.3 Virtual Filesystem and Persistence", 2)
    add_body(doc, "fs.js models directories, regular files, FIFOs, permissions, modification times, copy/move/remove operations, and path normalisation. The tree is serialised to localStorage after changes. Files distributed by the teaching staff are read by /api/terminal-files and merged into the learner's home directory.")

    add_heading(doc, "7.4 C and Process Simulation", 2)
    add_body(doc, "minic.js implements the subset of C required by the supplied practicals, including standard input/output, arrays, loops, functions, file APIs, process calls, signals, pipes, and System V message queues. process-table.js tracks PID, PPID, running/sleeping/zombie state, parent-child relationships, reaping, and orphan adoption by PID 1.")

    add_heading(doc, "7.5 Visual Concept Lessons", 2)
    add_body(doc, "ProcessConcept.jsx defines progressive lessons with labelled stages and selectable scenarios. Students can move one step at a time or play the animation, connecting source-code concepts with memory, queues, CPU scheduling, process tables, system reapers, and IPC objects.")

    add_callout(doc, "Design choice", "Routine practicals remain fast and local in the cached educational shell; advanced commands are deliberately delegated to a full WebVM environment. This hybrid approach balances classroom reliability with access to a broader Ubuntu userland.")

    add_page_break(doc)

    # Page 7: usage.
    add_heading(doc, "8. How to Use the Website", 1)
    steps = [
        "Open the website in a modern browser on the laboratory or personal device.",
        "Click inside the terminal and type help to view the available classroom commands.",
        "Use pwd, ls, cd, mkdir, cat, grep, chmod, and related commands exactly as in a Linux practical.",
        "Create or edit scripts and source files using nano filename, vi filename, vim filename, or gedit filename.",
        "Compile a C program using gcc program.c -o program and execute it with ./program.",
        "Choose Split Terminal when two coordinated shells are useful; both panes immediately see the same files.",
        "Use the OS Visual Lab routes for process, fork, zombie/orphan, and IPC demonstrations.",
        "Run ubuntu-vm when a task needs apt, python3, networking, or the full Ubuntu toolchain.",
        "Use Hard Reset only when cached files or running simulations must be cleared; shared lab data also resets automatically on a new browser day.",
    ]
    for text in steps:
        add_list_item(doc, text, usage_decimal_id)

    add_heading(doc, "8.1 Example Practical Workflow", 2)
    add_code_block(doc, "$ pwd\n/home/student\n$ ls\nhello.c  hello.sh  practicals  ...\n$ gcc hello.c -o hello\n$ ./hello\nHello, Ubuntu!")
    add_body(doc, "The screenshots on the following pages were captured from the production build of this repository on 2 September 2026.", after=0)

    add_page_break(doc)

    # Pages 8-10: screenshots.
    add_heading(doc, "9. Website Screenshots and Feature Demonstration", 1)
    add_body(doc, "Figure 1 shows a complete terminal workflow: location check, directory listing, C compilation, and execution. The visible Hello, Ubuntu! output verifies that the generated executable runs inside the classroom environment.")
    add_picture(doc, "terminal_compilation.png", 5.75, "Ubuntu Terminal Online showing pwd, ls, gcc compilation, and Hello Ubuntu program output")
    add_caption(doc, "Figure 1. Ubuntu-style terminal with filesystem access and C compilation.")
    add_picture(doc, "split_terminal_shared_fs.png", 5.75, "Two Ubuntu terminal panes using the same virtual filesystem and executing the same compiled program")
    add_caption(doc, "Figure 2. Split-terminal mode; both panes can access the shared executable and filesystem.")

    add_page_break(doc)
    add_heading(doc, "9.1 Interactive Operating Systems Visualisations", 2)
    add_body(doc, "The visual lab complements command execution with staged explanations. Learners can select a stage directly, move with Next, or play the sequence as an animation.")
    add_picture(doc, "process_visualizer.png", 5.75, "Operating Systems visual lab showing a process being scheduled onto the CPU")
    add_caption(doc, "Figure 3. Program-to-process lesson at the scheduling and CPU execution stage.")
    add_picture(doc, "zombie_process_visualizer.png", 5.75, "Zombie process lesson showing the parent collecting a child's stored exit status")
    add_caption(doc, "Figure 4. Zombie-process scenario showing wait(), the process table, and reaping.")

    add_page_break(doc)
    add_heading(doc, "9.2 Inter-Process Communication", 2)
    add_body(doc, "The communication lesson presents named pipes, message queues, shared memory, and signals using a consistent five-stage sequence: separate processes, create IPC, send, transfer, and receive. The selected screenshot shows a message travelling through a kernel-managed queue while the two process address spaces remain separate.")
    add_picture(doc, "ipc_message_queue.png", 6.1, "IPC visualisation showing a message transferred from process A to process B through a message queue")
    add_caption(doc, "Figure 5. System V-style message queue visualisation during data transfer.")
    add_callout(doc, "Teaching value", "The same concepts can first be observed visually and then practised through C programs that use fork(), wait(), pipes, files, or message-queue functions.")

    add_page_break(doc)

    # Page 11: testing and results.
    add_heading(doc, "10. Testing and Results", 1)
    add_body(doc, "The project was verified using an optimised production build and representative classroom tasks. The build completed successfully and generated routes for the terminal, file API, and all four visual lessons.")
    add_info_table(
        doc,
        ["Test", "Procedure", "Observed Result"],
        [
            ("Production build", "npm run build", "Compiled successfully; application routes generated"),
            ("Navigation", "pwd and ls", "Displayed /home/student and the expected cached/preloaded files"),
            ("C workflow", "gcc hello.c -o hello; ./hello", "Program executed and printed Hello, Ubuntu!"),
            ("Shared storage", "Run the same executable in both split panes", "Both terminals accessed the same file and output"),
            ("Content delivery", "Load /api/terminal-files", "Practical source files appeared under /home/student"),
            ("Process lesson", "Advance to the Run stage", "Scheduler, CPU, memory, and queue state updated visually"),
            ("Zombie lesson", "Select Zombie and advance to wait()", "Exit status and reaping relationship were displayed"),
            ("IPC lesson", "Select Message queue and Transfer", "Message packet moved through the queue from A toward B"),
        ],
        [1900, 3350, 4110],
        font_size=8.8,
    )
    add_callout(doc, "Result", "The implemented features satisfy the core requirement: students can complete common Ubuntu-oriented practical work on a Windows laboratory computer through the browser, while advanced Linux operations remain available through the full-VM option.")

    add_heading(doc, "10.1 Quality and Safety Characteristics", 2)
    for text in [
        "Routine commands execute locally in the browser and do not run arbitrary programs on the application server.",
        "Student-created files remain in that browser's storage unless reset; the design avoids a shared server-side student filesystem.",
        "Teacher-provided practical files are limited to readable UTF-8 text content.",
        "Hard Reset provides a visible recovery path for corrupted files or stuck simulations on shared machines.",
    ]:
        add_list_item(doc, text, bullet_id)

    add_page_break(doc)

    # Page 12: impact, limits, future, conclusion.
    add_heading(doc, "11. Benefits and Educational Impact", 1)
    for lead, detail in [
        ("Higher access. ", "Students are no longer limited to the subset of laboratory systems with a working Ubuntu installation."),
        ("Faster class start. ", "No dual-boot configuration or per-student virtual-machine setup is needed for routine exercises."),
        ("Practice continuity. ", "Browser-cached files survive refreshes, enabling learners to continue code and shell work during the same period."),
        ("Better conceptual understanding. ", "Animations connect commands and C system calls to CPU scheduling, process states, wait semantics, and IPC."),
        ("Simpler teaching distribution. ", "New or corrected practical files can be placed in the public terminal-files folder and delivered to learners on page load."),
    ]:
        add_list_item(doc, lead + detail, bullet_id, bold_lead=lead)

    add_heading(doc, "12. Limitations and Future Scope", 1)
    add_body(doc, "The cached shell is an educational implementation rather than a complete Linux kernel and GNU userland. Some commands implement the options needed for classroom use rather than every option available on Ubuntu. Browser localStorage has a capacity limit, binary files are not distributed through the practical-file endpoint, and the full WebVM option depends on internet access and the availability of webvm.io.")
    add_body(doc, "Future work can add authenticated student workspaces, instructor dashboards, cloud backup, assignment submission, automated output checking, more command options, additional C library functions, mobile-specific controls, accessibility testing, offline packaging, and institution-hosted full Linux containers or WebAssembly sandboxes.")

    add_heading(doc, "13. Conclusion", 1)
    add_body(doc, "Ubuntu Terminal Online directly addresses a real college laboratory constraint: many computers run Windows and some systems cannot reliably provide Ubuntu. The project converts any suitable browser into a practical Linux learning surface with a familiar terminal, persistent files, shell scripting, C compilation, process simulation, supplied laboratory programs, split-terminal collaboration, and interactive OS explanations.")
    add_body(doc, "Developed by Faiz Shaikh in his combined role as Teaching Assistant and first-year M.Tech student, the website is both an infrastructure solution and a teaching aid. It reduces dependence on machine configuration while giving students more opportunities to practise, experiment, and understand operating-system behaviour.")

    add_page_break(doc)

    # Page 13: appendix.
    add_heading(doc, "Appendix A - Supported Command Groups", 1)
    add_body(doc, "The following groups summarise the commands implemented in the cached classroom shell. Availability of every GNU option is not implied; the implementation focuses on the options used in laboratory exercises.")
    add_info_table(
        doc,
        ["Group", "Examples"],
        [
            ("Navigation and files", "pwd, cd, ls, tree, mkdir, rmdir, rm, touch, cp, mv, ln, stat, file, find, du, df"),
            ("Text processing", "cat, head, tail, wc, grep, sort, uniq, cut, tr, sed, awk, tee"),
            ("Permissions and FIFOs", "chmod, mkfifo"),
            ("Shell and scripting", "echo, printf, read, test, [, [[, export, env, alias, source, bash, sh, loops, functions, case"),
            ("System information", "date, cal, whoami, id, hostname, uname, uptime, free, ps, top"),
            ("Editors", "nano, vi, vim, gedit"),
            ("Toolchain", "gcc, cc, g++, make"),
            ("Session and recovery", "history, which, man, clear, reset, factory-reset, exit, logout"),
            ("Full environment", "ubuntu-vm for apt, apt-get, python3, networking, and broader Ubuntu tools"),
        ],
        [2200, 7160],
        font_size=9.2,
    )

    add_heading(doc, "Appendix B - Main Project Modules", 1)
    add_info_table(
        doc,
        ["Module", "Responsibility"],
        [
            ("app/page.jsx", "Main terminal workspace, split view, reset actions, and feature cards"),
            ("src/components/UbuntuTerminal.jsx", "Interactive terminal user interface and session integration"),
            ("src/lib/shell/interpreter.js", "Bash-like parsing, expansion, control flow, pipelines, and redirection"),
            ("src/lib/shell/commands.js", "Built-in Linux commands, editors, compiler commands, and VM launch"),
            ("src/lib/shell/fs.js", "Virtual filesystem, permissions, persistence, daily reset, and public-file merge"),
            ("src/lib/shell/minic.js", "Classroom C parsing/execution and operating-system APIs"),
            ("src/lib/shell/process-table.js", "PID/PPID state, zombie reaping, and orphan adoption"),
            ("src/components/process-concepts/", "Interactive process, fork, zombie/orphan, and IPC lessons"),
        ],
        [3500, 5860],
        font_size=8.8,
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Prepared from the verified project repository and production screenshots - 2 September 2026")
    set_run_font(r, size=8.8, color=MUTED, italic=True)

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build_document()
