from docx import Document


document = Document(r"C:\Users\Lenovo\Downloads\Shell_Scripting_Basic_Questions.docx")

for index, paragraph in enumerate(document.paragraphs):
    print(f"P{index}: {paragraph.text}")

for table_index, table in enumerate(document.tables):
    for row_index, row in enumerate(table.rows):
        cells = " | ".join(cell.text for cell in row.cells)
        print(f"T{table_index}R{row_index}: {cells}")
